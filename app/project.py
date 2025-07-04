# ====================================================================================
# Author: @dabihdev
# Year:   2024
# Python version: 3.6.5
# ====================================================================================

import os                                     # operations within the folders (works only on Windows)
from docx import Document                     # .docx file reading and writing
from bs4 import BeautifulSoup                 # XML/HTML parsing
from settings import *                        # import global settings
from datetime import datetime, timedelta      # date and time processing
import json                                   # JSON data parsing

class Project:
    """Class initializing a forecast project."""
    def __init__(self,
                 selected_day: int,
                 author_string: str
                 ):
        """Initialize the project."""

        # set forecast day
        self.forecast_day = datetime.today() + timedelta(days=selected_day)

        # set forecast date = project directory name = project name
        self.forecast_date = self.forecast_day.strftime('%d-%m-%Y') # convert to string format
        
        # get forecast date weekday
        self.forecast_weekday = weekdays_dict[self.forecast_day.weekday()]

        # project path
        self.path = output_dir+self.forecast_date+"/"                                                               
        
        # author(s) name(s)
        self.author_string = author_string

        # dictionary of file names inside the project directory (html is created only when exporting text to html page)
        self.filenames = {
            "author": self.author_string,
            "forecast_day": selected_day,
            "docx": f"previsione_{self.forecast_date}.docx",
            "svg": f"{self.forecast_date}.svg",
            "html": "",
            "json": f"{self.forecast_date}.json",
        }     

        # create project folder if not existing already
        if not os.path.exists(self.path):
            os.mkdir(self.path)
        
    def add_document(self):
        """Add a template .docx document formatted for the PRETEMP forecast to the project directory."""

        # Create empty docx file
        new_document = Document()

        # Add section: "TESTO BREVE"
        paragraph1 = new_document.add_paragraph()
        paragraph1.add_run("TESTO BREVE").bold = True # make it bold
        paragraph2 = new_document.add_paragraph("Inserire qui un riassunto della previsione.")

        # Add empty line
        new_document.add_paragraph()

        # Add header: "DISCUSSIONE"
        paragraph3 = new_document.add_paragraph()
        paragraph3.add_run("DISCUSSIONE").bold = True

        # Add section: "--sezione 1--"
        paragraph4 = new_document.add_paragraph()
        paragraph4.add_run("--sezione 1--").bold = True
        paragraph5 = new_document.add_paragraph("Inserire qui la discussione per questa sezione.")

        # Add section: "--sezione 2--"
        paragraph6 = new_document.add_paragraph()
        paragraph6.add_run("--sezione 2--").bold = True
        paragraph7 = new_document.add_paragraph("Inserire qui la discussione per questa sezione.")

        # Save the newly created docx document in the folder ./previsioni/ with the name formatted as "previsione_<today's date>.docx"
        new_document.save(self.path+self.filenames["docx"])

        # Log user
        print("> File .docx generato con successo!")

    
    def add_map(self) -> bool:
        """Create new SVG file in the forecast folder with template map and updated date and author."""

        # try opening the file "./assets/mappa.svg" and read xml content
        try:
            svg = open("../assets/mappa.svg", "rt")
        except FileNotFoundError:
            print("> File mappa.svg non trovato. Assicurarsi che la cartella assets sia presente nella cartella di preforema,\ne che contenga il file mappa.svg")
            return False # stop function and flag the program
        else:
            xml_content = svg.read()
            svg.close()            

        xml_soup = BeautifulSoup(xml_content, "xml")

        # Create updated string that will substitute the text
        newstring_date = f"Valida dalle ore 00:00 UTC alle 24:00 UTC di {self.forecast_weekday.lower()} {self.forecast_date} - Emessa: {todays_weekday.lower()} {todays_date} alle ore 15:00 UTC "
        newstring_author = "AUTORE: " + self.author_string

        # Find and update text object "DATA E AUTORE"
        xml_soup.find("tspan", {"id": "tspan25"}).string = newstring_date    # update dates
        xml_soup.find("tspan", {"id": "tspan26"}).string = newstring_author  # update author's name
        new_xml_content = str(xml_soup)                                      # store updated xml content as string

        
        # if file already exists overwrite it, otherwise create it
        with open(self.path+self.filenames["svg"], "w") as svg:
            svg.write(new_xml_content)

        # log user if .svg was successfully created
        print(f"> File .svg generato con successo!")
        
        # flag the program
        return True


    def add_html(self) -> bool:
        """Add html file in project directory, with updated date, time and author."""
        
        # try opening and reading the file "./assets/dd_mm_yyyy.html"
        try:
            html = open("../assets/dd_mm_yyyy.html", "rt")
        except FileNotFoundError:
            print("> File dd_mm_yyyy.html non trovato. Assicurarsi che la cartella assets sia presente nella cartella di preforema,\ne che contenga il file dd_mm_yyyy.html")
            return False # stop the function and flag the program
        else:
            html_soup = BeautifulSoup(html, "html5lib")
            html.close()
                   

        # Create updated title 
        weekday = self.forecast_weekday
        day = self.forecast_day.day
        month = months_dict[self.forecast_day.month]
        year = self.forecast_day.year
        new_title = f"PREVISIONE PER {weekday} {day} {month} {year}"

        # Create updated map URL
        map_url = f"https://www.pretemp.altervista.org/archivio/{year}/{month.lower()}/cartine/{self.forecast_date}.svg"
        
        # Create updated forecast time range
        new_forecast_time_range = f"Valida dalle ore 00:00 alle 24:00 UTC di {weekday.lower()} {day} {month.lower()} {year}"
        
        # Create updated authors string
        new_authors = "Previsore: " + self.author_string

        # Create date of forecast issue
        weekday = weekdays_dict[today.weekday()]
        day = today.day
        month = months_dict[today.month]
        year = today.year
        utc_hour = today.utcnow().strftime("%H:%M")
        new_forecast_issue_date = f"Emessa {weekday.lower()} {day} {month.lower()} {year} alle ore {utc_hour} UTC"

        # Update html elements
        html_soup.find("title", {"id": "window-title"}).string = new_title.lower() # window title
        html_soup.find("strong", {"id": "title-date"}).string = new_title # title
        html_soup.find("span", {"id": "forecast-time-range"}).string = new_forecast_time_range # forecast time range
        html_soup.find("img", {"id": "map-png"})["src"] = map_url # map URL
        html_soup.find("p", {"id": "issue-date"}).string = new_forecast_issue_date # forecast issue date
        html_soup.find("p", {"id": "authors"}).string = new_authors # forecast authors names

        # Save edited html in new_folder_path
        new_html_name = f'{self.forecast_date}.html'
        with open(self.path+new_html_name, "w") as html:
            html.write(str(html_soup.prettify(formatter="html")))

        # Log user
        print(f"> File .html generato con successo!")
        
        # Add file name to project dictionary of file names
        self.filenames["html"] = new_html_name

        # Flag the program if file .html was successfully generated
        return True


    def export_text_to_html(self):
        """Export forecast text from docx document to html page."""
        
        # Read docx paragraphs
        document = Document(self.path+self.filenames["docx"])
        paragraphs = []
        for paragraph in document.paragraphs:
            paragraphs.append(paragraph.text)
        
        # Read html code
        with open(self.path+self.filenames["html"]) as file:
            html_code = file.readlines()

        # Insert paragraphs in the html code. Make headers bold
        # 75 is the line in the html code before which the new lines should be inserted
        for i, paragraph in enumerate(paragraphs):
            if paragraph.startswith("-"): # --sections--, font: 12pt
                html_code.insert(75+i, f"<p class='testo' style='text-align: justify;'><span style='font-size: 12pt;'><strong>{paragraph}</strong></span></p>")
            elif paragraph.isupper(): # HEADERS, font: 14pt
                html_code.insert(75+i, f"<p class='testo' style='text-align: justify;'><span style='font-size: 14pt;'><strong>{paragraph}</strong></span></p>")
            else: # normal paragraphs, font: 12 pt
                html_code.insert(75+i, f"<p class='testo' style='text-align: justify;'><span style='font-size: 12pt;'>{paragraph}</span></p>")
        
        # Overwrite html document
        with open(self.path+self.filenames["html"], "w", encoding="utf-8") as output:
            output.writelines(html_code)


    def save_project_data(self):
        """Save project data to JSON file."""  
        output_file = open(self.path+self.forecast_date+".json", "w")
        json.dump(self.filenames, output_file, indent=2)
        output_file.close()
