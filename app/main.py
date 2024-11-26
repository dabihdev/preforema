# ====================================================================================
# Author: @dabihdev
# Year:   2024
# Python version: 3.6.5
# ====================================================================================


# IMPORT LIBRARIES ===================================================================
from bs4 import BeautifulSoup                 # XML parsing
from datetime import datetime, timedelta      # date and time processing
from docx import Document                     # .docx file reading and writing
import os                                     # operations within the folders (works only on Windows)
from settings import *                        # import global settings
# ====================================================================================

# FUNCTIONS ==========================================================================
def update_forecast_day(selected_day: int) -> tuple:
    """Set the current forecast day."""
    forecast_day = datetime.today() + timedelta(days=selected_day)
    forecast_day_date = forecast_day.strftime('%d-%m-%Y') # convert to string format
    forecast_day_weekday = weekdays_dict[forecast_day.weekday()]

    return forecast_day, forecast_day_date, forecast_day_weekday


def update_commands_description(selected_day: int) -> dict:
    """Update commands description with current forecast day."""
    commands["s"] =  f"selezionare il giorno di previsione (attuale: +{selected_day})."
    return commands


def display_commands():
    """Display to the user the available commands for this program."""
    print("")
    for key in commands.keys():
        print(f"- Digita [{key}] per {commands[key]}")


def print_readme():
    """Print the content of README.txt"""

    # read and print text from file
    with open("../README.txt", newline="", encoding="utf-8") as file:
        for line in file.readlines():
            print(line.strip())


def get_author_name() -> str:
    """Ask the user to insert his/her family name. Return formatted author name."""

    # Initialize output_string
    output_string = ""

    # Get user input
    authors_names = input("Specificare il cognome dell'autore. In caso di 2 o più autori, specificare i diversi cognomi separandoli con uno spazio> ")

    # Split authors names if more than one is given
    authors_names_list = authors_names.split() # split the names and put them in a list

    for name in authors_names_list:
        output_string += name.upper() + "/"    # attach the names in upper case to the output_string

    output_string = output_string[:-1]         # delete last '/'

    # return output string
    return output_string


def create_output_folder(output_dir=output_dir):
    """If output folder does not exist, create one."""
    if not os.path.exists(output_dir):
        os.mkdir(output_dir)


def create_new_forecast_folder(output_dir=output_dir) -> str:
    """
        Generate new folder with the forecast's date inside previsioni/. The newly
        generated folder will contain the forecast map and text for that day. The function
        also returns the path of the newly created folder.
    """

    # Do not generate folder if it already exists
    new_folder_path = output_dir+forecast_day_date + "/"
    if not os.path.exists(new_folder_path):
        os.mkdir(new_folder_path)

    return new_folder_path


def create_forecast_template_document(new_folder_path: str) -> str:
    """ Create a template .docx document formatted for the PRETEMP forecast and store it in
        the corresponding forecast folder. Also returns the name of the newly created file.
    """

    # Create empty docx file
    new_document = Document()

    # Add section: "TESTO BREVE"
    paragraph1 = new_document.add_paragraph()
    paragraph1.add_run("TESTO BREVE").bold = True # make it bold
    paragraph2 = new_document.add_paragraph("Inserire qui un riassunto della previsione.")

    # Add empty line
    new_document.add_paragraph()

    # Add header: "DISCUSSIONE"
    paragraph3 = new_document.add_paragraph()
    paragraph3.add_run("DISCUSSIONE").bold = True

    # Add section: "--sezione 1--"
    paragraph4 = new_document.add_paragraph()
    paragraph4.add_run("--sezione 1--").bold = True
    paragraph5 = new_document.add_paragraph("Inserire qui la discussione per questa sezione.")

    # Add section: "--sezione 2--"
    paragraph6 = new_document.add_paragraph()
    paragraph6.add_run("--sezione 2--").bold = True
    paragraph7 = new_document.add_paragraph("Inserire qui la discussione per questa sezione.")

    # Save the newly created docx document in the folder ./previsioni/ with the name formatted as "previsione_<today's date>.docx"
    output_name = "previsione_"+forecast_day_date+".docx"
    new_document.save(new_folder_path+output_name)

    return output_name
    

def create_updated_forecast_map(authors_string: str, new_folder_path: str) -> str:
    """
        Create new SVG file in the forecast folder with updated DATA E AUTORE.
        Return name of the new SVG file.
    """

    # Open the file "./assets/mappa.svg" and read xml content
    with open("../assets/mappa.svg", "rt") as svg:
        xml_content = svg.read()

    xml_soup = BeautifulSoup(xml_content, "xml")

    # Create updated string that will substitute the text
    newstring_date = f"Valida dalle ore 00:00 UTC alle 24:00 UTC di {forecast_day_weekday.lower()} {forecast_day_date} - Emessa: {todays_weekday.lower()} {todays_date} alle ore 15:00 UTC "
    newstring_author = "AUTORE: " + authors_string

    # Find and update text object "DATA E AUTORE"
    xml_soup.find("tspan", {"id": "tspan25"}).string = newstring_date    # update dates
    xml_soup.find("tspan", {"id": "tspan26"}).string = newstring_author  # update author's name
    new_xml_content = str(xml_soup)                                      # store updated xml content as string

    # Create new SVG file in documents and fill it with new_xml_content
    new_svg_name = f"mappa_{forecast_day_date}.svg" # set name of new svg file
    
    # if file already exists overwrite it, otherwise create it
    if os.path.isfile(new_folder_path+new_svg_name):
        with open(new_folder_path+new_svg_name, "w") as svg:
            svg.write(new_xml_content)
    else:
        with open(new_folder_path+new_svg_name, "x") as svg:
            svg.write(new_xml_content)

    return new_svg_name


def create_html_page(new_folder_path: str, new_authors: str) -> str:
    """
        Create new html file in new_folder_path, with updated date, time and author.
    """
    # Read html file content
    with open('../assets/dd_mm_yyyy.html') as html:
        html_soup = BeautifulSoup(html, "html5lib")

    # Create updated title 
    weekday = weekdays_dict[forecast_day.weekday()]
    day = forecast_day.day
    month = months_dict[forecast_day.month]
    year = forecast_day.year
    new_title = f"PREVISIONE PER {weekday} {day} {month} {year}"

    # Create updated forecast time range
    new_forecast_time_range = f"Valida dalle ore 00:00 alle 24:00 UTC di {weekday.lower()} {day} {month.lower()} {year}"

    # Create updated authors string
    new_authors = "Previsore: " + new_authors

    # Create date of forecast issue
    weekday = weekdays_dict[today.weekday()]
    day = today.day
    month = months_dict[today.month]
    year = today.year
    utc_hour = today.utcnow().strftime("%H:%M")
    new_forecast_issue_date = f"Emessa {weekday.lower()} {day} {month.lower()} {year} alle ore {utc_hour} UTC"

    # Update html elements
    html_soup.find("title", {"id": "window-title"}).string = new_title.lower() # window title
    html_soup.find("strong", {"id": "title-date"}).string = new_title # title
    html_soup.find("span", {"id": "forecast-time-range"}).string = new_forecast_time_range # forecast time range
    html_soup.find("p", {"id": "issue-date"}).string = new_forecast_issue_date # forecast issue date
    html_soup.find("p", {"id": "authors"}).string = new_authors # forecast authors names

    # Save edited html in new_folder_path
    new_html_name = f'{forecast_day_date}.html'
    if os.path.isfile(new_folder_path+new_html_name):
        with open(new_folder_path+new_html_name, "w") as html:
            html.write(str(html_soup.prettify(formatter="html")))
    else:
        with open(new_folder_path+new_html_name, "x") as html:
            html.write(str(html_soup.prettify(formatter="html")))

    return new_html_name


def write_on_html(new_folder_path: str, new_document_name: str, new_html_name: str):
    """Write forecast text on html page"""
    # Read docx paragraphs
    document = Document(new_folder_path+new_document_name)
    paragraphs = []
    for paragraph in document.paragraphs:
        paragraphs.append(paragraph.text)

    # Split paragraphs in 2 blocks: summary ("testo breve") and discussion ("discussione")
    # summary
    summary=[]
    for paragraph in paragraphs:
        if paragraph=="TESTO BREVE":
            continue
        elif paragraph=="DISCUSSIONE":
            break
        else:
            summary.append(paragraph)

    # discussion
    paragraphs.remove("TESTO BREVE")
    paragraphs.remove("DISCUSSIONE")
    for paragraph in summary:
        paragraphs.remove(paragraph)
    
    discussion = paragraphs[:]

    # Read html code
    with open(new_folder_path+new_html_name) as file:
        html_code = file.readlines()

    # Add discussion
    for i in range(len(discussion)):
        if discussion[i].startswith("-"): # make section headers bold
            html_code.insert(89+i, f"<p class='testo' style='text-align: justify;'><span style='font-size: 12pt;'><strong>{discussion[i]}</strong></span></p>")
        else:
            html_code.insert(89+i, f"<p class='testo' style='text-align: justify;'><span style='font-size: 12pt;'>{discussion[i]}</span></p>")
    
    # Add summary
    for i in range(len(summary)):
        if summary[i].startswith("-"): # make section headers bold
            html_code.insert(82+i, f"<p class='testo' style='text-align: justify;'><span style='font-size: 12pt;'><strong>{summary[i]}</strong></span></p>")
        else:
            html_code.insert(82+i, f"<p class='testo' style='text-align: justify;'><span style='font-size: 12pt;'>{summary[i]}</span></p>")

    # Overwrite html document
    with open(new_folder_path+new_html_name, "w", encoding="utf-8") as output:
        output.writelines(html_code)
# ====================================================================================

# MAIN ===============================================================================
if __name__ == "__main__":
    
    # splash screen
    print("\n========== PREFOREMA (PREtemp FOREcast MAnager) ==========")

    while(True):

        # update forecast day (default: tomorrow)
        forecast_day, forecast_day_date, forecast_day_weekday = update_forecast_day(selected_day)

        # update commands description
        commands = update_commands_description(selected_day)
        
        # prompt user
        display_commands()
        user_choice = input("> ")
        
         # select forecast day
        if user_choice == "s":
            print()
            day = input("Digitare il giorno di previsione come un numero intero (1 per domani, 2 per dopodomani, ecc.)> ") # prompt user
            try:
                day = int(day)
            except :
                print()
                print ('Il valore inserito non è numerico!')
            else:
                selected_day = day

        # initialize forecast project
        elif user_choice == "p":
            print()
            author_string = get_author_name()                                           # ask user input, return author's names
            print("Sto creando il progetto, attendere...")
            create_output_folder()                                                      # create output folder if there is none
            new_folder_path = create_new_forecast_folder()                              # create new forecast directory with the forecast's date 
            new_svg_name = create_updated_forecast_map(author_string, new_folder_path)  # create new SVG file inside forecast directory
            new_document_name = create_forecast_template_document(new_folder_path)      # create template document inside forecast directory
            os.startfile(new_folder_path+new_document_name)                       # open the newly created docx file in MS Word
            os.startfile(new_folder_path+new_svg_name)                            # open newly created svg map
            print("Progetto creato con successo!")
        
        # insert text from docx into html page
        elif user_choice == "e":
            try:
                print("Sto esportando il testo sulla pagina HTML...")
                new_html_name = create_html_page(new_folder_path, author_string)            # create template html page inside forecast directory
                write_on_html(new_folder_path, new_document_name, new_html_name)            # export forecast text into HTML page
            except:
                print("ERRORE: File di progetto non trovati. Prima di esportare il testo assicurarsi di aver creato un progetto.") # print error message
            else:
                print("Testo esportato!")
                os.startfile(new_folder_path+new_html_name)                           # open newly created HTML page
        
        # display info
        elif user_choice == "i":
            print()
            print_readme() # print README.txt
            print()
        
        # exit program
        elif user_choice == "x":
            print()
            print("Chiudo il programma. Ciao!")
            print()
            print("============== PREFOREMA, @DABIHDEV (2024) ===============")
            print()
            break # stop the program

        # handle exceptions
        else:
            print()
            print("Comando non riconosciuto, riprovare.")
            print()
# ====================================================================================
